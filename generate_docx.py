from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_viva_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('NimbusFS: Distributed File System Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Viva Preparation Guide')
    subtitle_format = subtitle.paragraph_format
    subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', level=1)
    p1 = doc.add_paragraph(
        'NimbusFS is a modern distributed file system simulation designed to demonstrate the fundamental principles of distributed computing. '
        'It features a centralized dashboard that manages file storage across multiple simulated nodes, ensuring high availability through data replication and consistency via a robust locking mechanism.'
    )
    
    # 2. Core Objectives
    doc.add_heading('2. Core Objectives', level=1)
    doc.add_paragraph('Data Redundancy: Automatically replicates every uploaded file across three storage nodes.', style='List Bullet')
    doc.add_paragraph('Data Integrity: Implements file locking to prevent unauthorized or concurrent modifications.', style='List Bullet')
    doc.add_paragraph('Monitoring & Transparency: Provides real-time statistics and activity logs.', style='List Bullet')
    
    # 3. System Architecture
    doc.add_heading('3. System Architecture', level=1)
    doc.add_paragraph('The project follows a Client-Coordinator-Storage model:')
    doc.add_paragraph('Frontend (Client): Built with React, it provides an intuitive interface for users.', style='List Bullet')
    doc.add_paragraph('Backend (Coordinator): A FastAPI server that handles all logic, metadata, and logging.', style='List Bullet')
    doc.add_paragraph('Storage Nodes: Three separate directories (node1, node2, node3) within the storage/ folder simulate independent physical storage servers.', style='List Bullet')
    
    # 4. Technical Stack
    doc.add_heading('4. Technical Stack', level=1)
    doc.add_heading('Backend:', level=2)
    doc.add_paragraph('FastAPI: High performance and native support for asynchronous programming.', style='List Bullet')
    doc.add_paragraph('SQLAlchemy: ORM to interact with the SQLite database.', style='List Bullet')
    
    doc.add_heading('Frontend:', level=2)
    doc.add_paragraph('React (Vite): For a fast, responsive user interface.', style='List Bullet')
    doc.add_paragraph('Tailwind CSS: For a sleek, professional aesthetic.', style='List Bullet')
    doc.add_paragraph('Framer Motion: Adds smooth animations.', style='List Bullet')
    
    # 5. Key Implementation Details
    doc.add_heading('5. Key Implementation Details', level=1)
    doc.add_heading('A. Replication Logic', level=2)
    doc.add_paragraph(
        'When a file is uploaded, the distribute_file function in file_manager.py loops through all available storage nodes and writes the file content to each. '
        'This ensures that even if two nodes fail, the data is still safe in the third.'
    )
    
    doc.add_heading('B. Concurrency & Locking', level=2)
    doc.add_paragraph(
        'Files can be Locked through the UI. This sets a locked boolean flag in the database. '
        'Any attempt to delete a locked file is rejected by the backend, demonstrating a race-condition prevention strategy.'
    )
    
    # 6. Expected Viva Questions & Answers
    doc.add_heading('6. Expected Viva Questions & Answers', level=1)
    
    q1 = doc.add_paragraph()
    q1.add_run('Q: Why did you choose FastAPI over Flask or Django?').bold = True
    doc.add_paragraph('A: FastAPI is significantly faster due to its asynchronous nature. It also provides automatic Swagger documentation, making it easier to test API endpoints during development.')
    
    q2 = doc.add_paragraph()
    q2.add_run('Q: How does your system handle node failure?').bold = True
    doc.add_paragraph('A: In this simulation, each file is replicated across all nodes. If a node were to go offline, the system would still have access to the file data in the remaining nodes, ensuring high availability.')
    
    q3 = doc.add_paragraph()
    q3.add_run('Q: What is the purpose of the Activity Log?').bold = True
    doc.add_paragraph('A: It provides an audit trail for every action taken in the system. This is crucial for debugging, security auditing, and monitoring user behavior.')
    
    doc.add_page_break()
    
    # Footer info
    footer = doc.add_paragraph('\n\nAuthor: Ayush\nProject: NimbusFS Distributed File System')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save the document
    filename = 'NimbusFS_Viva_Guide.docx'
    doc.save(filename)
    print(f'Successfully created {filename}')

if __name__ == '__main__':
    create_viva_guide()
